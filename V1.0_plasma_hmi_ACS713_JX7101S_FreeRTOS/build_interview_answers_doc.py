from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "低温等离子体消融控制项目面试问题回答总结_第2第3题更新版.docx"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        set_east_asia_font(r1)
        r2 = p.add_run(text[len(bold_prefix):])
        set_east_asia_font(r2)
    else:
        r = p.add_run(text)
        set_east_asia_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_east_asia_font(r)


def add_question(doc, idx, question, answer, key_points=None):
    h = doc.add_heading(f"{idx}. {question}", level=1)
    for run in h.runs:
        set_east_asia_font(run)

    add_para(doc, "建议回答：", bold_prefix="建议回答：")
    for paragraph in answer:
        add_para(doc, paragraph)

    if key_points:
        add_para(doc, "回答关键词：", bold_prefix="回答关键词：")
        add_bullets(doc, key_points)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for style_name, size, color, before, after in [
    ("Heading 1", 15, "2E74B5", 14, 7),
    ("Heading 2", 13, "2E74B5", 10, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("低温等离子体消融控制项目面试问题回答总结")
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor.from_string("0B2545")
set_east_asia_font(r)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("基于 STM32F103、FreeRTOS、互补 PWM、HMI 串口屏与电流保护")
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor.from_string("555555")
set_east_asia_font(r)

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(10)
r = intro.add_run(
    "使用建议：面试回答时先说明项目目标，再讲自己负责的模块，最后补充安全性、实时性和可优化方向。下面每个问题都按可直接口述的方式整理。"
)
set_east_asia_font(r)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.autofit = False
hdr = table.rows[0].cells
for i, text in enumerate(["主题", "面试关注点", "回答侧重点"]):
    hdr[i].text = text
    set_cell_shading(hdr[i], "F2F4F7")
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True
            set_east_asia_font(run)

rows = [
    ("项目介绍", "你是否真正理解系统架构", "业务目标、模块划分、个人工作"),
    ("PWM/互补 PWM", "是否理解功率驱动和安全风险", "频率、占空比、死区、MOE、安全停机"),
    ("FreeRTOS", "是否会做任务同步和实时保护", "事件标志组、互斥锁、信号量、消息队列"),
    ("HMI 协议", "是否能独立设计通信协议", "串口中断、状态机、帧格式、队列解耦"),
    ("难点与优化", "工程思维是否完整", "实时性、安全性、可靠性、后续改进"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        for p in cells[i].paragraphs:
            for run in p.runs:
                set_east_asia_font(run)
set_cell_margins(table)

add_question(
    doc,
    1,
    "请简单介绍一下你在这个项目中实现的东西",
    [
        "我这个项目是基于 STM32F103 和 FreeRTOS 做的低温等离子体消融设备控制系统，主要负责底层控制和人机交互部分。",
        "核心功能包括串口屏 HMI 通信、PWM 输出控制、单极/双极模式切换、电流采集与过流保护、脚踏开关控制、定脉冲输出，以及看门狗任务监控。",
        "软件上我把功能拆成几个 FreeRTOS 任务：电流采集任务负责 ADC + DMA 采样、RMS 电流计算和过流保护；控制任务负责脚踏和脉冲输出；UI 任务负责串口屏刷新；看门狗任务通过事件标志组判断各任务是否正常运行。",
    ],
    ["STM32F103", "FreeRTOS 多任务", "HMI 串口屏", "PWM 输出", "电流保护", "看门狗"],
)

add_question(
    doc,
    2,
    "在项目中为什么要用 PWM",
    [
        "因为这个系统需要产生高频高压来激发等离子体，但 MCU 本身不能直接输出高压，也不能直接产生等离子体。MCU 能输出的是 3.3V 左右的控制信号，所以需要通过 TIM1 产生 PWM 来控制后级功率开关器件。",
        "PWM 在这里不是直接输出能量，而是用来控制功率开关器件，让后级半桥或全桥电路产生高频交流或高频脉冲能量。可以理解为：MCU PWM 先经过驱动芯片，再控制 MOSFET 或 IGBT 半桥/全桥，把直流电变成高频方波，再经过后级升压、谐振或输出电路，最终形成用于激发等离子体的高频能量。",
        "在我的项目里，PWM 主要用于驱动后级半桥或全桥电路。通过修改 TIM1 的 ARR 改变频率，通过 CCR 改变占空比。HMI 下发频率或占空比后，我会调用对应的 PWM 设置函数动态更新参数，从而调节后级输出能量。",
    ],
    ["PWM 控制功率开关", "驱动芯片", "半桥/全桥", "高频方波", "ARR 控频率", "CCR 控占空比", "能量调节"],
)

add_question(
    doc,
    3,
    "为什么要实现互补 PWM 输出，有什么需要注意的地方",
    [
        "互补 PWM 主要是为了驱动半桥或全桥功率变换电路。互补信号控制上桥臂和下桥臂交替导通，让后级电路产生高频交流或高频脉冲能量。",
        "这里最需要注意的是死区时间。全桥或半桥里面，上桥臂和下桥臂不能同时导通；如果同时导通，就相当于电源正负极被短接，会发生直通，严重时会烧毁 IGBT 或 MOSFET。高频高压电源资料里也提到过，前期测试中如果不存在死区时间，可能会出现 IGBT 炸管。",
        "所以我们使用互补 PWM，不只是为了得到相反相位的两路信号，更重要的是配合死区时间，保证上下桥臂在切换瞬间不会同时导通，避免功率级损坏。",
        "我用的是 STM32 的高级定时器 TIM1，它支持主通道和互补通道输出。启动时除了启动 PWM 通道和互补通道，还要开启主输出 MOE；停止时要先关闭 MOE，再停止各通道，并把相关 GPIO 拉到安全电平。",
        "单极和双极模式下使用的输出通道不同，所以在单极模式下我会关闭部分通道，并手动配置引脚电平，避免悬空或误输出。",
    ],
    ["互补 PWM", "半桥/全桥", "死区", "防直通", "IGBT/MOSFET 保护", "TIM1 高级定时器", "MOE", "安全停机"],
)

add_question(
    doc,
    4,
    "为什么在其中用了一个事件标志组",
    [
        "我这里事件标志组主要用于任务级看门狗，也就是判断几个关键任务是否还活着。",
        "电流任务、控制任务、UI 任务分别设置不同的 bit。看门狗任务等待这些 bit 全部置位，如果在设定时间内都收齐，说明系统任务正常，就刷新硬件独立看门狗；如果某个任务卡死，没有打卡，就不刷新 IWDG，让系统复位。",
        "这样比单纯在一个地方喂狗更安全，因为它能防止主循环还在跑、但某个关键任务已经卡死的情况。",
    ],
    ["任务级看门狗", "多任务打卡", "等待所有 bit", "IWDG 刷新条件"],
)

add_question(
    doc,
    5,
    "如何在项目中用互斥锁，如何实现的",
    [
        "我用了互斥锁保护共享变量 current_display。这个变量由电流采集任务更新，同时 UI 任务要读取它并显示到串口屏，如果不加锁，可能出现读写竞争，导致显示值异常。",
        "实现方式是：电流任务计算出最新电流后，先申请 Mutex_Current，把值写入 current_display，然后马上释放。UI 任务刷新屏幕时，也先申请同一个互斥锁，把 current_display 拷贝到局部变量 safe_current，然后立即释放。",
        "这里的关键点是锁持有时间很短。我没有在持锁状态下做串口发送这种耗时操作，而是只在锁内完成一次变量拷贝，避免阻塞高优先级任务。",
    ],
    ["共享变量保护", "短临界区", "避免读写竞争", "不在锁内做串口发送"],
)

add_question(
    doc,
    6,
    "HMI 串口屏通信协议怎么做的",
    [
        "串口屏通过 USART2 和 MCU 通信。我实现了两部分：MCU 给屏幕发命令，比如更新文本、切换按钮状态；屏幕给 MCU 发控制指令，比如开关机、调频率、调占空比、发射固定脉冲数。",
        "接收端用了串口中断逐字节接收，并设计了一个简单状态机。协议大致是：帧头 0x55，命令类型，4 字节数据，最后是 0x0D 0x0A。命令类型 0x00 表示调频率，0x01 表示调占空比，0x02 表示触发脉冲输出。",
        "对于耗时操作，我没有在串口中断里直接执行，而是把脉冲数放入消息队列，由控制任务去处理，这样中断能尽快退出，系统实时性更好。",
    ],
    ["USART2", "中断接收", "状态机解析", "帧头帧尾", "消息队列解耦"],
)

add_question(
    doc,
    7,
    "实现这个项目时有什么难点",
    [
        "第一个难点是实时性和安全性的平衡。电流保护必须响应快，所以电流采集任务优先级较高，并且用 ADC + DMA 双缓冲，半满和全满中断通过信号量通知任务处理，避免 CPU 一直轮询。",
        "第二个难点是 PWM 输出的安全控制。互补 PWM 涉及半桥或全桥驱动，必须注意死区、MOE、停止顺序、引脚安全电平，尤其是过流时要快速关断 PWM。",
        "第三个难点是多任务之间的数据同步。电流数据既要用于保护，又要用于 UI 显示，还要避免 UI 串口发送影响控制实时性，所以分别用了互斥锁、信号量、消息队列、线程标志和事件标志组处理不同场景。",
    ],
    ["实时保护", "ADC + DMA 双缓冲", "PWM 安全关断", "多任务同步"],
)

add_question(
    doc,
    8,
    "后续有什么优化内容",
    [
        "后续可以增加更完善的故障状态机，比如过流、通信异常、传感器异常、脚踏异常分别记录错误码。",
        "HMI 协议可以增加 CRC 校验，避免串口干扰导致误指令。电流保护也可以做多级保护，比如硬件比较器瞬时保护加软件 RMS 保护。",
        "另外可以把用户参数存入 Flash，实现掉电保存；PWM 输出加入软启动，避免能量突变；再增加日志或调试接口，方便后期定位问题。",
    ],
    ["故障状态机", "CRC 校验", "硬件保护 + 软件保护", "Flash 参数保存", "软启动"],
)

add_question(
    doc,
    9,
    "什么是低温等离子体消融系统，主要实现什么功能，有什么用",
    [
        "低温等离子体消融系统是一种医疗或实验用能量控制系统。它通过高频电信号在电极附近激发介质形成等离子体，利用等离子体中的活性粒子对组织进行切割、消融或凝固。相比传统热切割，它的工作温度相对较低，对周围组织热损伤更小。",
        "我们这个项目主要实现的是设备控制部分，包括输出能量控制、单极/双极模式选择、脚踏触发、HMI 参数设置、电流检测和过流保护。",
        "简单说，就是让用户能通过串口屏设置参数，系统根据参数输出安全可控的高频 PWM 驱动信号，同时实时监测电流，出现异常时立即停止输出。",
    ],
    ["高频激发", "等离子体", "切割/消融/凝固", "低热损伤", "安全可控输出"],
)

add_question(
    doc,
    10,
    "项目中的脚踏开关控制指的是什么，是如何实现的",
    [
        "脚踏开关控制，指的是用户通过外部脚踏开关来控制设备是否输出 PWM，也就是实际设备中常见的“踩下开始输出，松开停止输出”。",
        "在我的项目里，脚踏开关接在 PA4，硬件上配置为上拉输入。未踩下时引脚为高电平，踩下时脚踏把引脚拉到低电平，所以软件通过读取 PA4 是否为 RESET 来判断脚踏是否被踩下。",
        "为了避免机械按键抖动，我在 FootPedal_IsPressed_HAL() 中做了多次采样确认，连续检测到低电平才认为脚踏真的踩下。在控制任务 StartControlTask() 中，只有 system_enable 为 1，并且当前没有定脉冲输出任务时，才允许脚踏控制 PWM。",
        "当检测到脚踏状态变化后，任务再延时 10ms 做一次确认。如果确认是踩下，就根据当前单极或双极档位调用 PWM_Start_HAL(mode)；如果确认是松开，就调用 PWM_Stop_HAL()。这样实现了踩下连续输出、松开立即停止的控制逻辑。",
    ],
    ["PA4 上拉输入", "踩下为低电平", "软件消抖", "边沿状态检测", "踩下启动 PWM", "松开停止 PWM"],
)

add_question(
    doc,
    11,
    "定脉冲输出指的是什么，是如何实现的",
    [
        "定脉冲输出，指的是用户通过 HMI 下发一个目标脉冲数，系统输出指定数量的 PWM 周期后自动停止，而不是像脚踏那样一直按住就持续输出。",
        "在我的项目里，HMI 串口协议中命令类型 0x02 表示发射固定脉冲数。串口中断解析到这个命令后，不直接在中断里执行耗时逻辑，而是把目标脉冲数放入 FreeRTOS 消息队列 PulseQueue。",
        "控制任务从 PulseQueue 中取出脉冲数后，会判断系统是否已经使能、脚踏是否未按下，然后调用 PulseControl_Start_HAL(received_pulse_counts, current_mode)。这个函数会保存目标脉冲数，清零当前计数，复位 TIM1 计数器，启动 PWM，并开启 TIM1 更新中断。",
        "之后每产生一个 PWM 周期，TIM1 会产生一次更新中断。在 PulseControl_TIM_IRQ_Handler() 中对 pulse_count 加 1；当 pulse_count 达到 pulse_target 时，自动调用 PulseControl_Stop_HAL() 停止 PWM。因为计数依赖 TIM1 更新事件，所以脉冲数量和 PWM 周期是同步的。",
    ],
    ["HMI 下发目标脉冲数", "0x02 命令", "消息队列解耦中断和任务", "TIM1 更新中断计数", "达到目标后自动停机"],
)

doc.add_section(WD_SECTION.CONTINUOUS)
closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(12)
r = closing.add_run("面试表达建议：")
r.bold = True
set_east_asia_font(r)
r = closing.add_run("每个问题控制在 30 到 90 秒内，先讲结论，再补一句项目里的具体实现，最后点到安全性或实时性，会更有工程味。")
set_east_asia_font(r)

doc.save(OUT)
print(OUT)
