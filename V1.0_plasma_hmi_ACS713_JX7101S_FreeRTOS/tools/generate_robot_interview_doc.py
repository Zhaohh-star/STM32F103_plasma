from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "智天成机器人嵌入式实习生面试速记问答.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 15, "1F4D78", 12, 5),
        ("Heading 2", 12.5, "2E74B5", 8, 3),
        ("Heading 3", 11, "1F4D78", 5, 2),
    ]:
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("智天成机器人嵌入式实习生面试速记问答")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("根据赵思琪简历项目整理 | 面试前快速复习版")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_fixed_layout(table)
    set_table_width(table, [9000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")
    p.add_run("  " + text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_qa(doc, q, a, followups=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Q: " + q)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("A: ").bold = True
    p.add_run(a)

    if followups:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("可能追问：")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("1F4D78")
        for item in followups:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Cm(0.85)
            bp.paragraph_format.space_after = Pt(1)
            bp.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_fixed_layout(table)
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_callout(
        doc,
        "面试主线",
        "把你的经历统一讲成：MCU 外设驱动 + RTOS 实时任务 + 通信协议 + 硬件联调 + 安全保护。这些能力可以自然迁移到机器人底层控制、传感器采集、执行器驱动和上位机通信。",
    )

    doc.add_heading("1. 一分钟自我介绍", level=1)
    p = doc.add_paragraph()
    p.add_run("建议版本：").bold = True
    p.add_run(
        "我是上海理工大学电子信息方向研二学生，主要做嵌入式软件开发。最近一个项目是基于 STM32F103 和 FreeRTOS 的高频电源控制系统，我负责 ADC/DMA 采样、电流 RMS 计算、PWM 控制、HMI 通信、Modbus 预留接口和任务调度。另一个项目是基于 CH582M 的 BLE 门禁系统，做过 BLE GATT 通信、RC522 RFID、SPI/GPIO 控制和事件调度。我比较熟悉 STM32 外设、FreeRTOS、通信协议和硬件联调，希望在机器人嵌入式岗位上接触电机、传感器和控制板开发。"
    )

    doc.add_heading("2. 简历项目与机器人岗位的对应关系", level=1)
    add_table(
        doc,
        ["简历经历", "你能强调的能力", "对应机器人场景"],
        [
            [
                "STM32F103 + FreeRTOS 高频电源控制系统",
                "ADC/DMA 实时采样、PWM 输出、Modbus/HMI 通信、任务调度、看门狗与故障保护",
                "电机驱动板、传感器采集板、执行器控制、安全急停、底层通信",
            ],
            [
                "CH582M RISC-V BLE 门禁系统",
                "BLE GATT、SPI 读卡、GPIO/PWM 控制、事件驱动、手机端指令交互",
                "机器人外设控制、无线配置、状态回传、低功耗/连接稳定性",
            ],
            [
                "医疗公司研发实习",
                "4G/WiFi 模块对接、设备状态上报、断点续传、PCBA 接口测试",
                "机器人联网、远程状态监控、产测调试、现场问题定位",
            ],
        ],
        [2350, 3400, 3250],
    )

    doc.add_heading("3. 高频电源项目高频追问", level=1)
    add_qa(
        doc,
        "为什么用 ADC + DMA？采样频率怎么考虑？",
        "ADC 直接中断读取会让 CPU 在高频采样下频繁进入中断，影响 FreeRTOS 任务调度。DMA 可以把采样数据自动搬到内存缓冲区，CPU 只在半传输或全传输完成时处理一批数据。采样频率要高于信号最高频率的两倍，同时结合 RMS 计算窗口、CPU 负载、DMA 缓冲长度和控制响应速度来定。",
        ["如果采样过快会有什么问题？", "DMA 半传输和全传输回调分别适合做什么？"],
    )
    add_qa(
        doc,
        "ACS712 电流 RMS 是怎么计算的？",
        "ACS712 输出带有零点偏置，通常先采集一段空载数据估计偏置，后续每个采样点减去偏置，再平方累加、求平均、开方得到 RMS。为了降低噪声影响，可以用滑动窗口或简单低通滤波。简历里可以强调你把浮点计算优化为整数运算，降低了 STM32F103 在高频采样下的运算压力。",
        ["偏置漂移怎么处理？", "RMS 窗口太长或太短分别有什么影响？"],
    )
    add_qa(
        doc,
        "PWM 为什么用 TIM1？如何保证输出安全？",
        "TIM1 是 STM32F103 的高级定时器，适合做较复杂的 PWM 输出，例如互补输出、死区控制和刹车/保护输入。安全上要做到默认关闭、异常时立即拉到安全态、互补输出避免上下桥直通、急停或故障时关断 PWM，并配合看门狗处理程序跑飞或任务卡死。",
        ["互补 PWM 和死区有什么作用？", "如果 GPIO 状态切换导致误触发，你怎么防？"],
    )
    add_qa(
        doc,
        "FreeRTOS 里任务怎么划分？",
        "可以按实时性划分：采样/控制任务优先级较高，负责处理 DMA 缓冲数据和更新控制状态；通信任务处理中低优先级的 HMI、Modbus、串口收发；UI 或日志任务优先级更低。任务间用队列传递数据，用信号量或事件标志通知状态变化，避免多个任务直接抢同一份资源。",
        ["队列、信号量、事件标志的区别？", "优先级设置不合理会造成什么现象？"],
    )
    add_qa(
        doc,
        "Modbus 读取 JX7101S 时如何处理异常？",
        "发送请求后要等待响应帧，检查地址、功能码、长度和 CRC。若超时、CRC 错误、异常功能码或长度不对，就丢弃该帧并记录错误，必要时重试。串口接收最好用状态机或空闲中断/DMA 方式处理，避免半包、粘包和错帧。",
        ["CRC 校验怎么做？", "RS485 半双工方向控制要注意什么？"],
    )

    doc.add_heading("4. CH582M BLE 门禁项目高频追问", level=1)
    add_qa(
        doc,
        "BLE 广播、连接、GATT 服务是什么关系？",
        "广播用于让手机发现设备，连接建立后手机作为 central 访问设备端 peripheral 暴露的 GATT 服务。服务下面有 characteristic，手机可以通过 write 下发命令，设备可以通过 notify 主动回传状态。",
        ["Notify 和 Indicate 有什么区别？", "配对和绑定有什么区别？"],
    )
    add_qa(
        doc,
        "TMOS 事件驱动和 FreeRTOS 有什么区别？",
        "TMOS 更像轻量事件调度，通常通过事件标志触发回调或任务片段，适合 BLE 协议栈这类事件型系统。FreeRTOS 则是完整 RTOS，有任务、优先级、阻塞、队列和同步机制，适合更复杂的并发任务管理。",
        ["事件处理里为什么不能长时间阻塞？", "BLE 连接稳定性会受哪些任务影响？"],
    )
    add_qa(
        doc,
        "RC522 刷卡流程怎么讲？",
        "RC522 通过 SPI 与 MCU 通信。典型流程是寻卡、防冲突、选卡、认证、读写块或读取 UID，然后把刷卡结果通过 BLE Notify 或本地状态机上报。工程上要处理重复刷卡、防抖、通信超时和异常卡片状态。",
        ["SPI 的 CPOL/CPHA 配错会怎样？", "如何判断是硬件接线问题还是协议问题？"],
    )

    doc.add_heading("5. 医疗公司实习可讲亮点", level=1)
    add_qa(
        doc,
        "实习经历怎么和机器人岗位关联？",
        "可以强调设备工程化能力：你做过通信模块接入、状态上报、异常恢复和 PCBA 接口测试。这些工作和机器人量产/调试很接近，机器人也需要联网、状态监控、接口验证、故障定位和现场问题复现。",
        ["4G/WiFi 模块通信不上你怎么查？", "PCBA 测试通常先查哪些接口？"],
    )

    doc.add_heading("6. 机器人嵌入式通用题速记", level=1)
    add_table(
        doc,
        ["问题", "回答要点"],
        [
            ["UART / SPI / I2C / RS485 / CAN 区别？", "UART 简单点对点；SPI 高速全双工但线多；I2C 线少多从机但速率有限；RS485 抗干扰远距离半双工；CAN 多主仲裁、强实时和可靠性，机器人电机/关节常见。"],
            ["机器人里为什么常用 CAN？", "抗干扰、支持多节点、带仲裁和错误检测，适合电机驱动器、传感器、控制板之间实时通信。"],
            ["电机不转怎么排查？", "先查电源和急停，再查使能/故障脚，再看 PWM/CAN/UART 指令，再查编码器/限位反馈，最后看驱动器保护和软件状态机。"],
            ["中断里能不能做耗时操作？", "不建议。中断里只做清标志、取数据、发通知，复杂计算放到任务里，否则会影响实时性并阻塞其他中断。"],
            ["看门狗怎么用？", "看门狗用于程序跑飞或关键任务卡死恢复。喂狗不能随便放主循环，最好由健康监测任务确认关键任务都正常后再喂。"],
            ["如何设计可靠串口协议？", "帧头、长度、命令字、序号、数据、CRC、超时、重试、状态机解析，必要时处理粘包/半包和版本兼容。"],
        ],
        [3000, 6000],
    )

    doc.add_heading("7. 最好准备的故障案例", level=1)
    add_numbered(
        doc,
        [
            "DMA Overrun 或采样数据异常：说明你如何用示波器/日志确认采样频率、DMA 回调、缓冲区长度和处理耗时。",
            "串口通信丢帧：说明你如何检查波特率、帧格式、CRC、接收状态机、RS485 方向控制和超时处理。",
            "PWM 输出异常或误触发：说明你如何确认定时器配置、GPIO 复用、默认安全态、互补输出和故障关断。",
            "BLE 连接不稳定：说明你如何减少事件处理耗时、调整连接参数、避免频繁 notify、检查手机端交互时序。",
        ],
    )

    doc.add_heading("8. 面试反问清单", level=1)
    add_bullets(
        doc,
        [
            "这个岗位主要负责机器人哪一层：电机驱动、传感器板、底盘控制、关节控制、通信网关还是整机联调？",
            "团队常用的 MCU/SoC、RTOS、通信总线是什么？是否大量使用 CAN、EtherCAT、RS485 或 ROS2？",
            "实习生入职后第一个月通常会参与什么任务：驱动开发、测试脚本、板级调试、协议适配还是问题定位？",
            "项目里对实时性和安全保护的要求是什么？例如急停、限位、故障上报、看门狗策略。",
            "如果我入职前想补齐知识，您建议我优先复习电机控制、CAN 通信、FreeRTOS 还是传感器融合？",
        ],
    )

    doc.add_heading("9. 面试时要避免的说法", level=1)
    add_table(
        doc,
        ["不要这样说", "更好的说法"],
        [
            ["我没做过机器人。", "我直接机器人经验不多，但做过 MCU、RTOS、通信、采样、PWM 和硬件联调，这些能迁移到底层控制。"],
            ["这个我不太会。", "这个我没有完整做过，但我理解它的基本原理；如果在项目里遇到，我会先从接口时序、协议和日志/波形定位。"],
            ["项目都是跟着做的。", "项目中我主要负责某几个模块，例如 ADC/DMA、PWM、Modbus 或 BLE GATT，并参与联调和异常处理。"],
        ],
        [3100, 5900],
    )

    doc.add_heading("10. 面试前 10 分钟速背", level=1)
    add_callout(
        doc,
        "核心句",
        "我不是只会写单个外设 Demo，我做过从采样、控制、通信、RTOS 任务到硬件联调和故障保护的完整嵌入式链路，这正是机器人底层开发需要的能力。",
    )
    add_bullets(
        doc,
        [
            "讲项目时固定结构：背景需求 -> 我的模块 -> 关键实现 -> 遇到问题 -> 如何验证。",
            "把高频电源项目讲成实时控制项目，把 BLE 门禁讲成无线外设与状态机项目，把实习讲成工程化调试项目。",
            "被问不会的内容，不要硬编。先说了解程度，再说排查思路和学习路径。",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("面试速记 | 机器人嵌入式实习生").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
